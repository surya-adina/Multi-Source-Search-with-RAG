import logging
import warnings
import time
import os
import re
import streamlit as st
from PyPDF2 import PdfReader
import pandas as pd
import docx
from langchain_community.document_loaders import RecursiveUrlLoader
from langchain_core.documents import Document
from bs4 import BeautifulSoup
from langchain_text_splitters import RecursiveCharacterTextSplitter
import google.generativeai as genai
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_groq import ChatGroq
from langchain_classic.chains.question_answering import load_qa_chain
from langchain_core.prompts import PromptTemplate
from dotenv import load_dotenv

# Suppress warnings
warnings.filterwarnings("ignore")

# Load environment variables
load_dotenv()

# Setup logging
LOG_FILE = "rag.log"
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s", handlers=[
    logging.FileHandler(LOG_FILE),
    logging.StreamHandler()
])
logger = logging.getLogger(__name__)


def simulate_typing(text, placeholder, typing_speed=0.0001):
    """
    Simulates typing animation for assistant response.

    Args:
        text (str): The full text of the response.
        placeholder (st.empty): Streamlit placeholder to update the text.
        typing_speed (float): Delay between displaying each character (in seconds).
    """
    typed_text = ""
    for char in text:
        typed_text += char
        placeholder.markdown(f"{typed_text}")
        time.sleep(typing_speed)
    placeholder.markdown(f"{text}")  # Ensure the full response is displayed



def is_common_greetings(query):
    patterns = [
        r"\b(hi|hello|hey|hiya|howdy|greetings|yo)\b", # Common greetings
        r"how (are|r) you\??", # Identity questions
        r"who (are|r) you\??", # Identity questions
        r"what('?s| is) your name\??", # Name questions
        r"what('?s| is) your role\??", # Role questions
        r"(hi|hello|hey|yo),? (who are you|what('?s| is) your name)\??", # Greeting + identity
        r"(hi|hello|hey|yo),? (what do you do|what can you do)\??", # Greeting + capability
        r"good (morning|afternoon|evening),? (who are you|what('?s| is) your role)\??", # Polite intros
        r"(what can you do for me)\??" # Informal assistance questions
    ]
    normalized = query.strip().lower()
    for pattern in patterns:
        if re.search(pattern, normalized):
            return True
    return False


# Function to extract text from PDFs
def get_pdf_text(doc):
    logger.info("Extracting text from uploaded Files...")
    text = ""
    # for doc in docs:
    pdf_reader = PdfReader(doc)
    for page in pdf_reader.pages:
        text += page.extract_text()
    logger.info("Doc text extraction complete.")
    return text


# Function to extract text from DOCX files
def get_docx_text(doc):
    logger.info("Extracting text from uploaded DOCX files...")
    text = ""
    # for doc in docs:
    docx_reader = docx.Document(doc)
    for para in docx_reader.paragraphs:
        text += para.text + "\n"
    logger.info("DOCX text extraction complete.")
    return text


# Function to extract text from CSV files
def get_csv_text(doc):
    logger.info("Extracting text from uploaded CSV files...")
    text = ""
    # for doc in docs:  # Iterate over multiple files if provided
    try:
        df = pd.read_csv(doc)  # Read CSV file
        text += df.to_string(index=False)  # Convert DataFrame to string
    except Exception as e:
        logger.error(f"Error reading CSV file {doc}: {str(e)}")
    logger.info("CSV text extraction complete.")
    return text


# Function to extract text from Excel files
def get_excel_text(doc):
    logger.info("Extracting text from uploaded Excel files...")
    text = ""
    # for doc in docs:
    try:
        df = pd.read_excel(doc, engine="openpyxl")  # Specify the engine
        text += df.to_string(index=False)  # Convert DataFrame to string
    except Exception as e:
        logger.error(f"Error reading Excel file {doc}: {str(e)}")
    logger.info("Excel text extraction complete.")
    return text


# Function to extract text from TXT files
def get_txt_text(doc):
    logger.info("Extracting text from uploaded TXT files...")
    text = ""
    # for doc in docs:
    text += doc.getvalue().decode("utf-8") + "\n"  # Decode bytes to string
    logger.info("TXT text extraction complete.")
    return text


def load_documents_from_web(CORPUS_SOURCE):

    loader = RecursiveUrlLoader(
        url=CORPUS_SOURCE,
        prevent_outside=True,
        base_url=CORPUS_SOURCE,
        max_depth=2,
        use_async=True
        # exclude_dirs=['https://www.csusb.edu/its/support/it-knowledge-base',
        #               'https://www.csusb.edu/its/support/knowledge-base']
        )
    raw_documents = loader.load()

    cleaned_documents = []
    for doc in raw_documents:
        cleaned_text = clean_text_from_html(doc.page_content)
        cleaned_documents.append(Document(page_content=cleaned_text, metadata=doc.metadata))

    return cleaned_documents


def clean_text_from_html(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')

    # Remove unnecessary elements
    for script_or_style in soup(['script', 'style', 'header', 'footer', 'nav']):
        script_or_style.decompose()

    main_content = soup.find('div', {'class': 'page-main-content'})
    #main_content = soup.find('main')
    if main_content:
        content = main_content.get_text(separator='\n')
    else:
        content = soup.get_text(separator='\n')

    return clean_text(content)


def clean_text(text):
    lines = (line.strip() for line in text.splitlines())
    cleaned_lines = [line for line in lines if line]
    return '\n'.join(cleaned_lines)


# Function to split text into chunks
def get_text_chunks(text):
    logger.info("Splitting text into chunks...")
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=150)
    chunks = text_splitter.split_text(text)
    logger.info(f"Total chunks created: {len(chunks)}")
    return chunks


# Embedding & index config
EMBED_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
INDEX_DIR = "faiss_index"

def get_embeddings():
    """Return a HuggingFaceEmbeddings instance (call where needed)."""
    return HuggingFaceEmbeddings(model_name=EMBED_MODEL)


# Function to create a FAISS vector store
def get_vector_store(text_chunks):
    """
    Build and save a FAISS vectorstore from text_chunks.
    Returns the vectorstore and embeddings instance.
    """
    logger.info("Creating FAISS vector store...")
    embeddings = get_embeddings()
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    vector_store.save_local(INDEX_DIR)
    logger.info("Vector store saved successfully.")
    return vector_store, embeddings


def load_vector_store():
    """
    Load the FAISS index if present. Returns (vectorstore, embeddings) or (None, None).
    """
    embeddings = get_embeddings()
    try:
        new_db = FAISS.load_local(INDEX_DIR, embeddings, allow_dangerous_deserialization=True)
        logger.info("Loaded FAISS index from disk.")
        return new_db, embeddings
    except Exception as e:
        logger.warning(f"Could not load FAISS index: {e}")
        return None, None



# Function to create the conversational chain
def get_conversational_chain():
    logger.info("Initializing conversational AI model...")
    prompt_template = """
    Answer the question as detailed as possible from the provided context. If the answer is not in the context, 
    just say, 'Answer is not available in the context'. Do not provide incorrect information.\n\n
    Context:\n {context}?\n
    Question: \n{question}\n
    Answer:
    """
    llm = ChatGroq(model="llama-3.1-8b-instant",temperature=0)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain = load_qa_chain(llm, chain_type="stuff", prompt=prompt)
    logger.info("Conversational AI model initialized.")
    return chain

def trim_docs(docs, max_chars=6000):
    """Trim combined doc text to stay under a rough prompt budget."""
    kept = []
    total = 0
    for d in docs:
        text = d.page_content or ""
        if total + len(text) > max_chars:
            remaining = max_chars - total
            if remaining <= 0:
                break
            d.page_content = text[:remaining]
        kept.append(d)
        total += len(d.page_content)
    return kept


# Function to handle user input
def user_input(user_question):
    logger.info(f"Processing user query: {user_question}")

    # Quick greeting handling
    if is_common_greetings(user_question):
        return f"Hi there! I am an AI Chat Bot, How can I help you today on the uploaded files...🤔!"

    try:
        # Load the vector store and embeddings
        new_db, embeddings = load_vector_store()

        if new_db is None:
            logger.warning("No FAISS index found. Please upload / process documents first.")
            return "I don't have any processed documents yet. Please upload files and click 'Submit & Process' first."

        # similarity_search will internally use the same embeddings for the query
        docs = new_db.similarity_search(user_question, k=2)
        docs = trim_docs(docs, max_chars=6000)

        if not docs:
            logger.warning("No relevant documents found in FAISS index.")
            return "No relevant documents found for your question."

        chain = get_conversational_chain()
        # adjust to whatever the chain expects; your chain returns a dict in your original code
        response = chain({"input_documents": docs, "question": user_question}, return_only_outputs=True)
        logger.info("Response generated successfully.")
        # Some chains return different keys; original used response["output_text"]
        # Defensive access:
        if isinstance(response, dict) and "output_text" in response:
            return response["output_text"]
        # fallback: convert to string
        return str(response)

    except Exception as e:
        logger.error(f"Error during query processing: {e}", exc_info=True)
        st.error("An error occurred while processing your request.")
        return "Sorry — something went wrong while processing your question."



# Main function for Streamlit app
def main():

    st.set_page_config("Multi Files Chatbot", page_icon=":scroll:")
    st.header("Multi-File's 🗐 - Chat Bot 🤖 ")

    # Display previous chat history
    # Initialize the session state for storing chat messages and chat history
    if "messages" not in st.session_state:
        st.session_state.messages = []  # Current chat messages
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []  # List of saved chats, each chat is a list of messages
    if "active_chat_index" not in st.session_state:
        st.session_state.active_chat_index = None  # Tracks which chat is currently active

    with st.sidebar:
        st.title("📁 Upload File's Section")

        is_data_submitted = False
        raw_text = ""
        files = st.file_uploader("Upload your Files(PDF, DOCX, CSV, XLSX, and TxT) & Click on Submit & Process", accept_multiple_files=True)

        if st.button("Submit & Process"):
            with st.spinner("Processing..."):
                # Variable to hold all the uploaded files text
                for file in files:
                    if file.name.endswith(".pdf"):
                        raw_text += get_pdf_text(file)
                    elif file.name.endswith(".docx"):
                        raw_text += get_docx_text(file)
                    elif file.name.endswith(".csv"):
                        raw_text += get_csv_text(file)
                    elif file.name.endswith(".xlsx") or file.name.endswith(".xls"):
                        raw_text += get_excel_text(file)
                    elif file.name.endswith(".txt"):
                        raw_text += get_txt_text(file)

                # raw_text = get_pdf_text(files)
                # print("Raw text:")
                # print(raw_text)
                is_data_submitted = True

        st.write("---")
        st.title("📁 URL Upload Section")
        corpus_source_url = st.text_input("Past you website URL & Click on Submit URL")

        if st.button("Submit URL"):
            with st.spinner("Processing..."):
                # Variable to hold all the uploaded URL text data
                data_in_object = load_documents_from_web(corpus_source_url)
                data = data_in_object[0]
                raw_text += data.metadata["title"] + " "
                raw_text += data.page_content
                is_data_submitted = True

        if is_data_submitted:
            with st.spinner("Processing..."):
                text_chunks = get_text_chunks(raw_text)
                get_vector_store(text_chunks)
                st.success("Processing complete!")
                logger.info("PDF processing and vector storage completed.")

        # Display previous chat history in the sidebar
        st.write("---")
        st.write("### Chat History 📝")
        if st.button("New Chat"):
            if st.session_state.messages:  # Save current chat if it has messages
                if st.session_state.active_chat_index is not None:
                    # Update the existing chat in history if the user was in a previous chat
                    st.session_state.chat_history[st.session_state.active_chat_index] = st.session_state.messages.copy()
                else:
                    # Save the new chat if it was not linked to an existing history
                    st.session_state.chat_history.append(st.session_state.messages.copy())
            st.session_state.messages = []  # Clear chat messages for new chat
            st.session_state.active_chat_index = None  # Reset active chat index
            st.rerun()  # <-- CHANGED FROM st.experimental_rerun() TO st.rerun()

        if st.session_state.chat_history:
            for i, chat in enumerate(st.session_state.chat_history):
                chat_name = f"Chat {i + 1} ({len(chat)} messages)"
                if st.button(chat_name, key=f"chat_{i}"):
                    st.session_state.messages = chat.copy()  # Load selected chat history
                    st.session_state.active_chat_index = i  # Set the active chat index
                    st.rerun()  # <-- CHANGED FROM st.experimental_rerun() TO st.rerun()

    # Display existing chat messages
    for message in st.session_state.messages:
        if message["role"] == "user":
            with st.chat_message("user"):
                st.write(message["content"])
        else:
            with st.chat_message("assistant"):
                st.write(message["content"])

    # Chat input field
    if prompt := st.chat_input("Ask a Question from the Data uploaded .. ⌨"):
        # Save the user's message
        st.session_state.messages.append({"role": "user", "content": prompt})

        # If the user is in an active chat, update the corresponding chat in history
        if st.session_state.active_chat_index is not None:
            st.session_state.chat_history[st.session_state.active_chat_index] = st.session_state.messages.copy()

        # Display user's message
        with st.chat_message("user"):
            st.write(prompt)

        # Simulated assistant response
        response = user_input(prompt)

        # Save assistant's response
        st.session_state.messages.append({"role": "assistant", "content": response})

        # If the user is in an active chat, update the corresponding chat in history
        if st.session_state.active_chat_index is not None:
            st.session_state.chat_history[st.session_state.active_chat_index] = st.session_state.messages.copy()

        # Display assistant's response with typing animation
        with st.chat_message("assistant"):
            # st.write(response)
            placeholder = st.empty()  # Placeholder for typing effect
            simulate_typing(response, placeholder, typing_speed=0.0001)


if __name__ == "__main__":
    main()